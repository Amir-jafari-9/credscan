from pathlib import Path
import xml.etree.ElementTree as ET
from .utils import detect_encoding
from typing import Iterator

def read_plain_text(file_path: Path) -> Iterator[str]:
    try:
        encoding = detect_encoding(file_path)
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            for line in f: yield line
    except Exception as e:
        yield f"ERROR: {str(e)}"

def read_xml(file_path: Path) -> Iterator[str]:
    try:
        tree = ET.parse(str(file_path))
        for elem in tree.iter():
            if elem.text and elem.text.strip(): yield elem.text
    except Exception:
        yield from read_plain_text(file_path)

def read_docx(file_path: Path) -> Iterator[str]:
    try:
        from docx import Document
        doc = Document(str(file_path))
        for para in doc.paragraphs:
            if para.text.strip(): yield para.text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip(): yield cell.text
    except Exception as e:
        yield f"ERROR: DOCX failed ({str(e)})"

def read_xlsx(file_path: Path) -> Iterator[str]:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(file_path), read_only=True, data_only=True)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None: yield str(cell)
    except Exception as e:
        yield f"ERROR: XLSX failed ({str(e)})"

READERS = {
    '.txt': read_plain_text,
    '.rtf': read_plain_text,
    '.xml': read_xml,
    '.hml': read_xml,
    '.docx': read_docx,
    '.xlsx': read_xlsx,
}

def get_reader(file_path: Path):
    return READERS.get(file_path.suffix.lower(), read_plain_text)
